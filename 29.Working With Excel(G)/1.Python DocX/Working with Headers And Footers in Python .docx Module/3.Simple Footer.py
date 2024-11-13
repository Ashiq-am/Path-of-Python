# Import docx NOT python-docx
import docx

# Create an instance of a word document
doc = docx.Document()

# Choosing the top most section of the page
section = doc.sections[0]

# Calling the footer
footer = section.footer

# Calling the paragraph already present in
# the footer section
footer_para = footer.paragraphs[0]

# Adding text in the footer
footer_para.text = "This is a footer..."

# Add a Title to the document
doc.add_heading('GeeksForGeeks', 0)

# Now save the document to a location
doc.save('gfg.docx')
