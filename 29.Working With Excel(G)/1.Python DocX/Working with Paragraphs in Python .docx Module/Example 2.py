# Import docx NOT python-docx
import docx

# Create an instance of a word document
doc = docx.Document()

# Add a Title to the document
doc.add_heading('GeeksForGeeks', 0)

# Adding multilined paragraph
doc.add_paragraph('''GeeksforGeeks is a Computer Science portal for geeks.
It contains well written, well thought and well explained computer science
and programming articles, quizzes etc.''')

# Now save the document to a location
doc.save('gfg.docx')
