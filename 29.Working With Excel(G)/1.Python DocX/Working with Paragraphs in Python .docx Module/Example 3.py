# Import docx NOT python-docx
import docx

# Create an instance of a word document
doc = docx.Document()

# Add a Title to the document
doc.add_heading('GeeksForGeeks', 0)

# Adding Normal Texted paragraph
doc.add_heading('Normal Style:', 3)
doc.add_paragraph('GeeksforGeeks is a Computer Science portal for geeks.',
				style = 'Normal')

# Adding Body Text Styled paragraph
doc.add_heading('Body Style:', 3)
doc.add_paragraph('GeeksforGeeks is a Computer Science portal for geeks.',
				style = 'Body Text')

# Adding Caption Styled paragraph
doc.add_heading('Caption Style:', 3)
doc.add_paragraph('GeeksforGeeks is a Computer Science portal for geeks.',
				style = 'Caption')

# Adding Title Styled paragraph
doc.add_heading('Title Style:', 3)
doc.add_paragraph('GeeksforGeeks is a Computer Science portal for geeks.',
				style = 'Title')

# Adding Heading Styled paragraph
doc.add_heading('Heading 1 Style:', 3)
doc.add_paragraph('GeeksforGeeks is a Computer Science portal for geeks.',
				style = 'Heading 1')

# Adding Macro Text Styled paragraph
doc.add_heading('Macro Text Style:', 3)
doc.add_paragraph('GeeksforGeeks is a Computer Science portal for geeks.',
				style = 'macro')

# Adding Quoted Style paragraph
doc.add_heading('Quoted Style:', 3)
doc.add_paragraph('GeeksforGeeks is a Computer Science portal for geeks.',
				style = 'Quote')

# Adding TOC Heading Styled paragraph
doc.add_heading('TOC Heading Style:', 3)
doc.add_paragraph('GeeksforGeeks is a Computer Science portal for geeks.',
				style = 'TOC Heading')

# Adding Subtitle Styled paragraph
doc.add_heading('Subtitle Style:', 3)
doc.add_paragraph('GeeksforGeeks is a Computer Science portal for geeks.',
				style = 'Subtitle')

# Adding No Spacing Styled paragraph
doc.add_heading('No Spacing Style:', 3)
doc.add_paragraph('GeeksforGeeks is a Computer Science portal for geeks.',
				style = 'No Spacing')

# Now save the document to a location
doc.save('gfg.docx')
