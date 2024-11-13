# Import docx NOT python-docx
import docx

# Create an instance of a word document
doc = docx.Document()

# Add a Title to the document
doc.add_heading('GeeksForGeeks', 0)

# Adding paragraph
para = doc.add_paragraph('GeeksforGeeks is a Computer Science portal for geeks.')

# Setting keep_with_next as True
para.keep_with_next = True

# Now save the document to a location
doc.save('gfg.docx')
