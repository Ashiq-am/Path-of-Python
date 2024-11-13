# Import docx NOT python-docx
import docx

# Create an instance of a word document
doc = docx.Document()

# Add a Title to the document
doc.add_heading('GeeksForGeeks', 0)

# Image in its native size
doc.add_heading('Image in Native Size:', 3)
doc.add_picture('logo.png')

# Now save the document to a location
doc.save('gfg.docx')
