# Import docx NOT python-docx
import docx
from docx.shared import Inches

# Create an instance of a word document
doc = docx.Document()

# Add a Title to the document
doc.add_heading('GeeksForGeeks', 0)


# Image with defined size
doc.add_heading('Image with Defined Size:', 3)
doc.add_picture('logo.png', width=Inches(2), height=Inches(2))

# Now save the document to a location
doc.save('gfg.docx')
