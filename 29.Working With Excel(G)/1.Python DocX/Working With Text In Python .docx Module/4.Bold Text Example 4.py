# Import docx NOT python-docx
import docx

# Create an instance of a word document
doc = docx.Document()

# Add a Title to the document
doc.add_heading('GeeksForGeeks', 0)

# Creating paragraph with some content
para = doc.add_paragraph(
	'''GeeksforGeeks is a Computer Science portal for geeks.''')

# Adding more content to paragraph and applying italics to true
para.add_run(
	''' It contains well written, well thought and well-explained ''').italic = True

# Adding more content to paragraph
para.add_run('''computer science and programming articles, quizzes etc.''')

# Now save the document to a location
doc.save('gfg.docx')
