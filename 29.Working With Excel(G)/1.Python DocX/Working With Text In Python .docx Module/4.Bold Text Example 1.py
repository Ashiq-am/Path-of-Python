# Import docx NOT python-docx
import docx

# Create an instance of a word document
doc = docx.Document()

# Add a Title to the document
doc.add_heading('GeeksForGeeks', 0)

# Creating paragraph
para = doc.add_paragraph()

# Adding content to paragraph
bold_para = para.add_run(
	'''GeeksforGeeks is a Computer Science portal for geeks. It contains well written, well thought and well-explained computer science and programming articles, quizzes etc.''')

# Setting bold to true
bold_para.bold = True

# Now save the document to a location
doc.save('gfg.docx')
