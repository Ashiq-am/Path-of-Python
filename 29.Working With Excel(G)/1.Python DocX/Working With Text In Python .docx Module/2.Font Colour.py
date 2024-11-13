# Import docx NOT python-docx
import docx
from docx.shared import RGBColor

# Create an instance of a word document
doc = docx.Document()

# Add a Title to the document
doc.add_heading('GeeksForGeeks', 0)

# Adding paragraph
doc.add_heading('Font Colour:', 3)
para = doc.add_paragraph().add_run(
	'GeeksforGeeks is a Computer Science portal for geeks.')

# Adding forest green colour to the text
# RGBColor(R, G, B)
para.font.color.rgb = RGBColor(0x22, 0x8b, 0x22)

# Now save the document to a location
doc.save('gfg.docx')
