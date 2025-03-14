# Import docx NOT python-docx
import docx

# Create an instance of a word document
doc = docx.Document()

# Add a Title to the document
doc.add_heading('GeeksForGeeks', 0)

# Adding paragraph with new font Style
doc.add_heading('Font Style: Roboto', 3)
para = doc.add_paragraph().add_run(
	'GeeksforGeeks is a Computer Science portal for geeks.')
# Setting new font style
para.font.name = 'Roboto'

# Adding paragraph with default font Style
doc.add_heading('Font Style: Default [Cambria]', 3)
doc.add_paragraph(
	'GeeksforGeeks is a Computer Science portal for geeks.')

# Now save the document to a location
doc.save('gfg.docx')
