# Import docx NOT python-docx
import docx

# Create an instance of a word document
doc = docx.Document()

# Add a heading of level 0 (Also called Title)
doc.add_heading('Title for the document', 0)

# Add a heading of level 1
doc.add_heading('Heading level 1', 1)

# Add a heading of level 2
doc.add_heading('Heading level 2', 2)

# Add a heading of level 3
doc.add_heading('Heading level 3', 3)

# Add a heading of level 4
doc.add_heading('Heading level 4', 4)

# Add a heading of level 5
doc.add_heading('Heading level 5', 5)

# Add a heading of level 6
doc.add_heading('Heading level 6', 6)

# Add a heading of level 7
doc.add_heading('Heading level 7', 7)

# Add a heading of level 8
doc.add_heading('Heading level 8', 8)

# Add a heading of level 9
doc.add_heading('Heading level 9', 9)

# Now save the document to a location
doc.save('gfg.docx')
