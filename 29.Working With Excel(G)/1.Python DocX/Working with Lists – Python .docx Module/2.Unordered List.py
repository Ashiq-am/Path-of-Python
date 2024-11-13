# Import docx NOT python-docx
import docx

# Create an instance of a word document
doc = docx.Document()

# Add a Title to the document
doc.add_heading('GeeksForGeeks', 0)

# Adding list of style name 'List Bullet'
doc.add_heading('Style: List Bullet', 3)
# Adding points to the list named 'List Number'
doc.add_paragraph('The first item in an unordered list.',
				style='List Bullet')

doc.add_paragraph('The second item in an unordered list.',
				style='List Bullet')

doc.add_paragraph('The third item in an unordered list.',
				style='List Bullet')

# Adding list of style name 'List Bullet 2'
doc.add_heading('Style: List Bullet 2', 3)
# Adding points to the list named 'List Number'
doc.add_paragraph('The first item in an unordered list.',
				style='List Bullet 2')

doc.add_paragraph('The second item in an unordered list.',
				style='List Bullet 2')

doc.add_paragraph('The third item in an unordered list.',
				style='List Bullet 2')

# Adding list of style name 'List Bullet 3'
doc.add_heading('Style: List Bullet 3', 3)
# Adding points to the list named 'List Number'
doc.add_paragraph('The first item in an unordered list.',
				style='List Bullet 3')

doc.add_paragraph('The second item in an unordered list.',
				style='List Bullet 3')

doc.add_paragraph('The third item in an unordered list.',
				style='List Bullet 3')

# Now save the document to a location
doc.save('gfg.docx')
